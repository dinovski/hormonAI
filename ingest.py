#!/usr/bin/env python
"""
ingest.py — type-aware, embedding-model-agnostic ingestion for hormonAI.

Handles two document shapes, which need different treatment:

  - FAQ   : atomic Q/A pairs. One chunk per question (question + answer kept
            together). Same behaviour as the original ingest_faq.py.
  - ARTICLE: free prose (.docx, .md, .txt, .pdf). Split on the document's own
            headings where available (docx styles / markdown '#'); PDFs, which
            carry no reliable heading structure, are treated as one flowing
            section. Long sections are split into fixed-size (~word count) child
            chunks with overlap, each keeping a heading breadcrumb and a link to
            its parent section text (for future parent-document retrieval).

Both are emitted into ONE unified per-item schema so the retriever stays
uniform:

    {
      "id":          "<source_id>:<n>",
      "source_id":   "<file stem>",
      "source_type": "faq" | "article",
      "lang":        "en" | "fr",
      "section":     "<display topic>",        # FAQ section or article top heading
      "heading_path":"<A > B > C>",            # breadcrumb
      "question":    "<FAQ question | article heading breadcrumb>",
      "answer":      "<FAQ answer | article child chunk text>",
      "parent_id":   "<section id> | None",
      "parent_text": "<full section text> | None",
      "q_paraphrases": [],
    }

The retriever (rag_core) reads question/section/answer/source_type, so this
schema is a compatible superset of the original one.

Embedding is model-agnostic: pass any SentenceTransformer with
--embedding-model. For models that need instruction prefixes (e.g. e5:
"query:" / "passage:") pass --passage-prefix / --query-prefix; the query
prefix is stored in the payload and used by rag_core at query time. mpnet and
BGE-M3 need no prefixes.

Outputs, per language, match the names rag_core loads:
    {data_dir}/{out_prefix}_{lang}_qa.pkl
    {data_dir}/{out_prefix}_{lang}_bm25.pkl
    {data_dir}/{out_prefix}_{lang}_index_q.faiss
    {data_dir}/{out_prefix}_{lang}_index_qa.faiss
    {data_dir}/{out_prefix}_{lang}_index_qp.faiss
"""
from __future__ import annotations

import os
import re
import json
import pickle
import argparse
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

# NOTE: heavy deps (faiss, sentence_transformers, numpy, rank_bm25) are imported
# lazily inside build_and_save() so the parsing/chunking logic can be imported
# and unit-tested without them.


# ---------------------------------------------------------------------------
# Text utilities
# ---------------------------------------------------------------------------

def normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def tokenize_for_bm25(text: str) -> List[str]:
    return re.findall(r"\b\w+\b", (text or "").lower())


def word_count(text: str) -> int:
    return len(re.findall(r"\S+", text or ""))


def split_sentences(text: str) -> List[str]:
    """Lightweight, language-agnostic sentence splitter (EN + FR)."""
    text = normalize_spaces(text)
    if not text:
        return []
    # Split after ., !, ? (optionally followed by closing quotes/brackets) + space.
    parts = re.split(r"(?<=[.!?])[\"')\]]*\s+", text)
    return [p.strip() for p in parts if p.strip()]


def chunk_text(text: str, chunk_size_words: int = 250, overlap_words: int = 40) -> List[str]:
    """
    Pack whole sentences into chunks of about `chunk_size_words`, with
    `overlap_words` of trailing context carried into the next chunk. Sentences
    are never split unless a single sentence exceeds the chunk size, in which
    case it is hard-split on word boundaries.
    """
    text = normalize_spaces(text)
    if not text:
        return []

    sentences = split_sentences(text)
    chunks: List[str] = []
    cur: List[str] = []
    cur_words = 0

    def flush():
        nonlocal cur, cur_words
        if cur:
            chunks.append(" ".join(cur).strip())
        cur = []
        cur_words = 0

    for sent in sentences:
        sw = word_count(sent)

        # A single oversized sentence: hard-split on words.
        if sw > chunk_size_words:
            flush()
            words = sent.split()
            for i in range(0, len(words), chunk_size_words):
                chunks.append(" ".join(words[i:i + chunk_size_words]).strip())
            continue

        if cur_words + sw > chunk_size_words and cur:
            # Close the current chunk, then seed the next with trailing overlap.
            flush()
            if overlap_words > 0 and chunks:
                tail_words = chunks[-1].split()[-overlap_words:]
                cur = [" ".join(tail_words)]
                cur_words = len(tail_words)
        cur.append(sent)
        cur_words += sw

    flush()
    return [c for c in chunks if c]


# ---------------------------------------------------------------------------
# Unified item
# ---------------------------------------------------------------------------

def make_item(
    source_id: str,
    source_type: str,
    lang: str,
    n: int,
    section: str,
    heading_path: str,
    question: str,
    answer: str,
    parent_id: Optional[str] = None,
    parent_text: Optional[str] = None,
) -> Dict[str, Any]:
    return {
        "id": f"{source_id}:{n}",
        "source_id": source_id,
        "source_type": source_type,
        "lang": lang,
        "section": normalize_spaces(section) or ("General" if lang == "en" else "Général"),
        "heading_path": normalize_spaces(heading_path),
        "question": normalize_spaces(question),
        "answer": (answer or "").strip(),
        "parent_id": parent_id,
        "parent_text": (parent_text or None),
        "q_paraphrases": [],
    }


# ---------------------------------------------------------------------------
# FAQ handler (reuses the original robust parser)
# ---------------------------------------------------------------------------

def parse_faq_docx(path: str, lang: str) -> List[Dict[str, Any]]:
    from docx import Document  # lazy
    import ingest_faq  # reuse the battle-tested Q/A parser

    doc = Document(path)
    raw = ingest_faq.parse_docx_into_qa(doc, lang)  # [{section, question, answer}]
    source_id = os.path.splitext(os.path.basename(path))[0]
    items: List[Dict[str, Any]] = []
    for n, it in enumerate(raw):
        items.append(make_item(
            source_id=source_id, source_type="faq", lang=lang, n=n,
            section=it["section"], heading_path=it["section"],
            question=it["question"], answer=it["answer"],
        ))
    return items


# ---------------------------------------------------------------------------
# Article handler
# ---------------------------------------------------------------------------

@dataclass
class Section:
    heading_path: List[str] = field(default_factory=list)
    paragraphs: List[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return normalize_spaces(" ".join(self.paragraphs))


def _sections_from_docx(path: str) -> List[Section]:
    from docx import Document  # lazy
    doc = Document(path)
    sections: List[Section] = []
    stack: List[str] = []          # heading text by depth index
    cur = Section(heading_path=[])

    def start_section(new_path: List[str]):
        nonlocal cur
        if cur.paragraphs:
            sections.append(cur)
        cur = Section(heading_path=list(new_path))

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = ""
        try:
            style = (p.style.name or "").lower()
        except Exception:
            style = ""
        m = re.match(r"heading\s+(\d+)", style)
        if m:
            level = int(m.group(1))
            stack = stack[: level - 1]
            while len(stack) < level - 1:
                stack.append("")
            stack.append(text)
            start_section(stack)
        else:
            cur.paragraphs.append(text)

    if cur.paragraphs:
        sections.append(cur)
    return sections


def _sections_from_markdown(path: str) -> List[Section]:
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    sections: List[Section] = []
    stack: List[str] = []
    cur = Section(heading_path=[])

    def start_section(new_path: List[str]):
        nonlocal cur
        if cur.paragraphs:
            sections.append(cur)
        cur = Section(heading_path=list(new_path))

    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if m:
            level = len(m.group(1))
            heading = m.group(2).strip()
            stack = stack[: level - 1]
            while len(stack) < level - 1:
                stack.append("")
            stack.append(heading)
            start_section(stack)
        elif line.strip():
            cur.paragraphs.append(line.strip())

    if cur.paragraphs:
        sections.append(cur)
    return sections


def extract_pdf_pages(path: str) -> Tuple[List[str], str]:
    """
    Return (per-page text, backend). PyMuPDF (fitz) is the default because it
    reconstructs reading order across multi-column layouts far better than
    pypdf; pypdf is the fallback if PyMuPDF is unavailable or errors.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = [(pg.get_text("text") or "") for pg in doc]
        doc.close()
        return pages, "pymupdf"
    except Exception:
        from pypdf import PdfReader
        reader = PdfReader(path)
        out: List[str] = []
        for pg in reader.pages:
            try:
                out.append(pg.extract_text() or "")
            except Exception:
                out.append("")
        return out, "pypdf"


_PAGE_NUM_RE = re.compile(r"^(page\s*)?\d+(\s*[/of]{1,2}\s*\d+)?$", re.IGNORECASE)


def clean_pdf_pages(pages: List[str]) -> str:
    """
    Clean common PDF extraction artifacts before chunking:
      - drop running headers/footers (short lines that repeat across many pages),
      - drop bare page-number lines ("12", "Page 3", "3 / 14"),
      - de-hyphenate words split across line wraps ("meno-\\npause" -> "menopause").
    Returns a single normalized text block.
    """
    from collections import Counter

    n = len(pages)
    per_page_lines: List[List[str]] = []
    line_counts: "Counter[str]" = Counter()
    for pg in pages:
        lines = [ln.strip() for ln in (pg or "").splitlines()]
        per_page_lines.append(lines)
        for ln in {l for l in lines if l}:
            line_counts[ln] += 1

    # A line is boilerplate if it is short and recurs on a large share of pages.
    # For very short documents (n < 4) the threshold is never reached, so nothing
    # is stripped -- avoids over-cleaning leaflets.
    recur_threshold = max(3, int(0.4 * n))
    boiler = {
        ln for ln, c in line_counts.items()
        if c >= recur_threshold and len(ln.split()) <= 12
    }

    kept_pages: List[str] = []
    for lines in per_page_lines:
        keep = []
        for ln in lines:
            if not ln or ln in boiler:
                continue
            if _PAGE_NUM_RE.match(ln):
                continue
            keep.append(ln)
        kept_pages.append("\n".join(keep))

    text = "\n".join(kept_pages)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)   # de-hyphenate line wraps
    return normalize_spaces(text)


def _sections_from_pdf(path: str) -> List[Section]:
    """
    Extract a PDF into one flowing section per document. PDFs carry no reliable
    heading structure (no styles), so rather than guess headings and risk
    fragmenting the text, the whole document is treated as a single section
    (after header/footer + hyphenation cleaning) and handed to the sentence-aware
    chunker. The heading breadcrumb falls back to the file name; the chunk text
    itself carries the meaning for retrieval.
    """
    pages, _backend = extract_pdf_pages(path)
    full = clean_pdf_pages(pages)
    if not full:
        return []
    return [Section(heading_path=[], paragraphs=[full])]


def parse_article(
    path: str,
    lang: str,
    chunk_size_words: int = 250,
    overlap_words: int = 40,
) -> List[Dict[str, Any]]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        sections = _sections_from_docx(path)
    elif ext in (".md", ".markdown", ".txt"):
        sections = _sections_from_markdown(path)
    elif ext == ".pdf":
        sections = _sections_from_pdf(path)
    else:
        raise ValueError(f"Unsupported article format: {ext} ({path})")

    source_id = os.path.splitext(os.path.basename(path))[0]
    items: List[Dict[str, Any]] = []
    n = 0
    for si, sec in enumerate(sections):
        breadcrumb = " > ".join([h for h in sec.heading_path if h]) or source_id
        top = next((h for h in sec.heading_path if h), source_id)
        section_text = sec.text
        if not section_text:
            continue
        parent_id = f"{source_id}#sec{si}"
        for chunk in chunk_text(section_text, chunk_size_words, overlap_words):
            items.append(make_item(
                source_id=source_id, source_type="article", lang=lang, n=n,
                section=top, heading_path=breadcrumb,
                # The "question"/topic side of the index is the heading breadcrumb;
                # the chunk text becomes the answer body.
                question=breadcrumb,
                answer=chunk,
                parent_id=parent_id,
                parent_text=section_text,
            ))
            n += 1
    return items


# ---------------------------------------------------------------------------
# Embedding text builders (shared by ingest and, in spirit, rag_core queries)
# ---------------------------------------------------------------------------

def _lead_snippet(text: str, n_words: int = 40) -> str:
    """First ~n_words of a chunk, used as an article's topic representation."""
    words = re.findall(r"\S+", text or "")
    return " ".join(words[:n_words])


def q_text_for(item: Dict[str, Any]) -> str:
    if item["source_type"] == "faq":
        return f"Section: {item['section']}\nQuestion: {item['question']}"
    # Articles have no question. Using only the heading breadcrumb makes the
    # Q-side degenerate (for PDFs it is just the file name), so all article
    # chunks look identical to the Q / Q+paraphrase channels. Use the heading
    # plus the chunk's leading text as a topic proxy so those channels carry
    # real content.
    return f"Section: {item['heading_path']}\n{_lead_snippet(item['answer'])}"


def qa_text_for(item: Dict[str, Any]) -> str:
    if item["source_type"] == "faq":
        return f"Section: {item['section']}\nQuestion: {item['question']}\nAnswer: {item['answer']}"
    return f"Section: {item['heading_path']}\n{item['answer']}"


def qp_text_for(item: Dict[str, Any]) -> str:
    paras = item.get("q_paraphrases") or []
    if paras:
        joined = "\n".join(f"- {p}" for p in paras)
        return q_text_for(item) + f"\nParaphrases:\n{joined}"
    return q_text_for(item)


def bm25_doc_for(item: Dict[str, Any]) -> List[str]:
    blob = f"{item['heading_path']} {item['section']} {item['question']} {item['answer']}"
    return tokenize_for_bm25(blob)


# ---------------------------------------------------------------------------
# Build + save (embedding-model agnostic)
# ---------------------------------------------------------------------------

def build_and_save(
    items: List[Dict[str, Any]],
    lang: str,
    data_dir: str,
    out_prefix: str,
    embedding_model: str,
    passage_prefix: str = "",
    query_prefix: str = "",
    manifest_meta: Optional[Dict[str, Any]] = None,
) -> None:
    import numpy as np
    import faiss
    from rank_bm25 import BM25Okapi
    from sentence_transformers import SentenceTransformer

    os.makedirs(data_dir, exist_ok=True)
    base = f"{out_prefix}_{lang}"
    qa_path = os.path.join(data_dir, f"{base}_qa.pkl")
    bm25_path = os.path.join(data_dir, f"{base}_bm25.pkl")
    idx_q_path = os.path.join(data_dir, f"{base}_index_q.faiss")
    idx_qa_path = os.path.join(data_dir, f"{base}_index_qa.faiss")
    idx_qp_path = os.path.join(data_dir, f"{base}_index_qp.faiss")

    print(f"[{lang}] {len(items)} items "
          f"(faq={sum(1 for i in items if i['source_type']=='faq')}, "
          f"article={sum(1 for i in items if i['source_type']=='article')})")

    q_texts = [passage_prefix + q_text_for(it) for it in items]
    qa_texts = [passage_prefix + qa_text_for(it) for it in items]
    qp_texts = [passage_prefix + qp_text_for(it) for it in items]
    bm25_docs = [bm25_doc_for(it) for it in items]

    print(f"[{lang}] loading embedding model: {embedding_model}")
    model = SentenceTransformer(embedding_model)

    def encode(texts: List[str]) -> "np.ndarray":
        emb = model.encode(texts, convert_to_numpy=True, show_progress_bar=True).astype("float32")
        faiss.normalize_L2(emb)
        return emb

    def build_index(emb: "np.ndarray") -> "faiss.Index":
        index = faiss.IndexFlatIP(emb.shape[1])
        index.add(emb)
        return index

    print(f"[{lang}] encoding q / qa / qp ...")
    faiss.write_index(build_index(encode(q_texts)), idx_q_path)
    faiss.write_index(build_index(encode(qa_texts)), idx_qa_path)
    faiss.write_index(build_index(encode(qp_texts)), idx_qp_path)

    bm25 = BM25Okapi(bm25_docs)

    with open(qa_path, "wb") as f:
        pickle.dump({
            "items": items,
            "embedding_model_name": embedding_model,
            "passage_prefix": passage_prefix,
            "query_prefix": query_prefix,
            "language": lang,
            "manifest_meta": manifest_meta or {},
            "schema_version": 2,
        }, f)

    with open(bm25_path, "wb") as f:
        pickle.dump({"bm25": bm25, "bm25_docs": bm25_docs, "language": lang}, f)

    print(f"[{lang}] wrote {qa_path}, {bm25_path}, and 3 FAISS indexes.")


# ---------------------------------------------------------------------------
# Sources: manifest or convenience flags
# ---------------------------------------------------------------------------

def load_sources(args: argparse.Namespace) -> List[Dict[str, str]]:
    sources: List[Dict[str, str]] = []
    if args.manifest:
        with open(args.manifest, "r", encoding="utf-8") as f:
            entries = json.load(f)
        for e in entries:
            sources.append({
                "path": e["path"],
                "type": e.get("type", "faq"),
                "lang": e.get("lang", args.language),
            })
    for p in (args.faq or []):
        sources.append({"path": p, "type": "faq", "lang": args.language})
    for p in (args.article or []):
        sources.append({"path": p, "type": "article", "lang": args.language})
    return sources


def collect_items(sources: List[Dict[str, str]], args: argparse.Namespace) -> Dict[str, List[Dict[str, Any]]]:
    by_lang: Dict[str, List[Dict[str, Any]]] = {}
    for src in sources:
        path, stype, lang = src["path"], src["type"], src["lang"]
        if not os.path.exists(path):
            raise FileNotFoundError(path)
        if stype == "faq":
            items = parse_faq_docx(path, lang)
        elif stype == "article":
            items = parse_article(path, lang, args.chunk_size, args.chunk_overlap)
        else:
            raise ValueError(f"Unknown source type: {stype}")
        print(f"  parsed {len(items):4d} items  [{stype:7s} {lang}]  {path}")
        by_lang.setdefault(lang, []).extend(items)

    # Reassign globally-unique ids per language (keeps positional idx stable).
    for lang, items in by_lang.items():
        for n, it in enumerate(items):
            it["idx"] = n
    return by_lang


def main() -> None:
    ap = argparse.ArgumentParser(description="Type-aware, model-agnostic ingestion for hormonAI.")
    ap.add_argument("--manifest", default=None,
                    help="JSON list of {path, type: faq|article, lang: en|fr}.")
    ap.add_argument("--faq", action="append", help="FAQ .docx path (repeatable).")
    ap.add_argument("--article", action="append", help="Article .docx/.md/.txt/.pdf path (repeatable).")
    ap.add_argument("--language", "-l", choices=["en", "fr"], default="en",
                    help="Language for --faq/--article convenience flags.")
    ap.add_argument("--data-dir", default="data")
    ap.add_argument("--out-prefix", default="kb",
                    help="Output basename; rag_core loads '<prefix>_<lang>_*' and '<prefix>_all_*'. Default: kb.")
    ap.add_argument("--embedding-model",
                    default="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
    ap.add_argument("--passage-prefix", default="",
                    help="Prefix prepended to document text before encoding (e.g. 'passage: ' for e5).")
    ap.add_argument("--query-prefix", default="",
                    help="Prefix rag_core prepends to the query (e.g. 'query: ' for e5). Stored in payload.")
    ap.add_argument("--chunk-size", type=int, default=250, help="Article chunk size in words.")
    ap.add_argument("--chunk-overlap", type=int, default=40, help="Article chunk overlap in words.")
    ap.add_argument("--shared", action=argparse.BooleanOptionalAction, default=True,
                    help="Also build a combined '<prefix>_all_*' index across all languages "
                         "(shared multilingual space, each item lang-tagged) for cross-lingual "
                         "fallback. Default: on. Disable with --no-shared.")
    ap.add_argument("--per-language", action=argparse.BooleanOptionalAction, default=True,
                    help="Also build the per-language '<prefix>_<lang>_*' indexes. Default: on.")
    args = ap.parse_args()

    sources = load_sources(args)
    if not sources:
        ap.error("No sources. Use --manifest, or --faq/--article with --language.")

    print("Collecting items...")
    by_lang = collect_items(sources, args)
    meta = {"chunk_size": args.chunk_size, "chunk_overlap": args.chunk_overlap}

    if args.per_language:
        for lang, items in by_lang.items():
            build_and_save(
                items=items, lang=lang, data_dir=args.data_dir, out_prefix=args.out_prefix,
                embedding_model=args.embedding_model,
                passage_prefix=args.passage_prefix, query_prefix=args.query_prefix,
                manifest_meta=meta,
            )

    # Combined shared-space index across every language. Items keep their own
    # per-item "lang" tag; the retriever prefers same-language and falls back
    # cross-lingual. Written as '<prefix>_all_*' (lang bucket name = "all").
    if args.shared:
        combined: List[Dict[str, Any]] = []
        for lang in sorted(by_lang):
            combined.extend(by_lang[lang])
        build_and_save(
            items=combined, lang="all", data_dir=args.data_dir, out_prefix=args.out_prefix,
            embedding_model=args.embedding_model,
            passage_prefix=args.passage_prefix, query_prefix=args.query_prefix,
            manifest_meta={**meta, "languages": sorted(by_lang)},
        )
    print("Done.")


if __name__ == "__main__":
    main()
